"""
Export government-services questions ranked by semantic similarity to a query.
Outputs a Word .docx file for manual review.
"""

import duckdb
import numpy as np
from sentence_transformers import SentenceTransformer
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pathlib

# ── Config ────────────────────────────────────────────────────────────────────
QUERY   = 'accessing government services after the fire'
TOP_N   = 200   # how many deduplicated results to include in the doc
DB_PATH = pathlib.Path(__file__).parent / 'data' / 'eaton_pipeline.duckdb'
OUT_DOC = pathlib.Path(__file__).parent / 'govt_similarity_ranked_v2.docx'

GOVT_TAGS = ['fema_ia', 'sba', 'dua', 'dmv', 'tax', 'd_snap',
             'debris', 'permits', 'insurance', 'utilities']

# ── Load data ─────────────────────────────────────────────────────────────────
print("Connecting to database…")
conn = duckdb.connect(str(DB_PATH), read_only=True)
tag_filter = ' OR '.join([f"list_contains(program_area_tags, '{t}')" for t in GOVT_TAGS])

df = conn.execute(f"""
    SELECT id, cluster_id, cluster_size, program_area_tags, question_type_tags,
           subreddit, score, created_utc::DATE as date,
           COALESCE(body_clean, body) as body
    FROM records
    WHERE is_question = TRUE AND ({tag_filter})
    ORDER BY id
""").df()
print(f"  {len(df)} government-services candidates loaded")

# ── Encode ─────────────────────────────────────────────────────────────────────
print("Loading sentence-transformer model…")
model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')

texts = df['body'].fillna('').tolist()
print(f"Encoding {len(texts)} candidates + query…")
all_embs = model.encode([QUERY] + texts, normalize_embeddings=True,
                        batch_size=128, show_progress_bar=True)
query_emb = all_embs[0]
cand_embs = all_embs[1:]
df['sim'] = cand_embs @ query_emb

# ── Deduplicate: best per cluster ──────────────────────────────────────────────
group_key = df['cluster_id'].where(df['cluster_id'] >= 0, other=df['id'].astype('int64', errors='ignore'))
# For singleton rows cluster_id == -1; use row index as unique key instead
df = df.copy()
df['_group'] = df.apply(
    lambda r: f"cluster_{int(r['cluster_id'])}" if r['cluster_id'] >= 0 else f"singleton_{r['id']}",
    axis=1
)

top = (
    df.sort_values('sim', ascending=False)
      .groupby('_group', sort=False)
      .first()
      .sort_values('sim', ascending=False)
      .head(TOP_N)
      .reset_index(drop=True)
)

print(f"Top {len(top)} deduplicated results (one per cluster) for: '{QUERY}'")

# ── Build Word doc ─────────────────────────────────────────────────────────────
doc = Document()

# Narrow margins for readability
for section in doc.sections:
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# Title
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("Eaton Fire — Government Services Questions")
run.bold = True
run.font.size = Pt(16)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_para.add_run(
    f'Ranked by semantic similarity to: "{QUERY}"\n'
    f'{len(top)} results (deduplicated; one exemplar per cluster)'
)
sub_run.font.size = Pt(10)
sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()  # spacer

for rank, (_, row) in enumerate(top.iterrows(), start=1):
    cid   = row['cluster_id']
    sz    = row['cluster_size']
    sim   = row['sim']
    pa    = row['program_area_tags']
    qt    = row['question_type_tags']
    tags  = ', '.join(pa)  if (pa  is not None and len(pa)  > 0) else '—'
    qtags = ', '.join(qt)  if (qt  is not None and len(qt)  > 0) else '—'

    cluster_label = (f"cluster {int(cid)} · {int(sz)} similar posts"
                     if cid >= 0 else "singleton")

    # Header line
    hdr = doc.add_paragraph()
    hdr.paragraph_format.space_before = Pt(10)
    hdr.paragraph_format.space_after  = Pt(2)

    r1 = hdr.add_run(f"#{rank}  ")
    r1.bold = True
    r1.font.size = Pt(11)

    r2 = hdr.add_run(f"sim={sim:.3f}  ·  {cluster_label}  ·  r/{row['subreddit']}  ·  score={row['score']}  ·  {row['date']}")
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

    # Tags line
    tag_para = doc.add_paragraph()
    tag_para.paragraph_format.space_before = Pt(0)
    tag_para.paragraph_format.space_after  = Pt(3)
    tr = tag_para.add_run(f"program: {tags}   |   type: {qtags}")
    tr.font.size = Pt(8)
    tr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    tr.italic = True

    # Body — full text, no truncation
    body_para = doc.add_paragraph()
    body_para.paragraph_format.space_before = Pt(0)
    body_para.paragraph_format.space_after  = Pt(4)
    br = body_para.add_run(row['body'] or '')
    br.font.size = Pt(10)

    # Thin rule
    sep = doc.add_paragraph('─' * 80)
    sep.paragraph_format.space_before = Pt(4)
    sep.paragraph_format.space_after  = Pt(2)
    sep.runs[0].font.size = Pt(7)
    sep.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

doc.save(str(OUT_DOC))
print(f"\n✓ Saved: {OUT_DOC}")
